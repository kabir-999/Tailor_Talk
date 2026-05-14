from __future__ import annotations

import csv
import hashlib
import json
import logging
import re
import shutil
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from PyPDF2 import PdfReader

from app.models.schemas import FileResult
from app.services.embeddings import semantic_scores

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".csv",
    ".xlsx",
    ".xls",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".gif",
    ".mp4",
    ".mov",
    ".m4v",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
VIDEO_EXTENSIONS = {".mp4", ".mov", ".m4v"}
UPLOAD_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx", ".xls", *IMAGE_EXTENSIONS, *VIDEO_EXTENSIONS}


@dataclass
class LocalDocument:
    path: Path
    content: str


@dataclass
class LocalQuery:
    terms: list[str]
    phrases: list[str]
    extensions: set[str]
    filenames: set[str]
    modified_after: datetime | None = None


class LocalSearchService:
    def __init__(self, root: Path, uploads_only: bool = True) -> None:
        self.root = root
        self.uploads_only = uploads_only
        self.cache_dir = self.root / ".search_cache"

    @property
    def configured(self) -> bool:
        return self.root.exists() and self.root.is_dir()

    @property
    def uploads_dir(self) -> Path:
        return self.root / "uploads"

    def search(self, query: str, limit: int = 10) -> tuple[list[FileResult], str]:
        if not self.configured:
            raise RuntimeError(f"Local folder does not exist: {self.root}")

        parsed = self._parse_query(query)
        paths = [path for path in self._iter_files() if self._passes_filters(path, parsed)]
        docs = [self._load_document(path, allow_ocr=False) for path in paths]
        docs = [doc for doc in docs if doc.content or doc.path.name]

        if not docs:
            scope = "uploaded files only" if self.uploads_only else "Assignment folder"
            return [], f"{scope}: no files matched the requested file type/date filters"

        ranked = self._rank_documents(docs, parsed)
        if self._should_use_ocr(query, parsed):
            existing = {doc.path for _score, doc in ranked}
            needs_ocr = len(ranked) < min(limit, 3)
            matched_ocr_paths = {path for path in existing if self._ocr_candidate(path)}
            unmatched_ocr_paths = {path for path in paths if path not in existing and self._ocr_candidate(path)}
            ocr_paths = sorted(matched_ocr_paths | (unmatched_ocr_paths if needs_ocr else set()))
            if ocr_paths:
                ocr_docs = [self._load_document(path, allow_ocr=True) for path in ocr_paths]
                ocr_doc_paths = {doc.path for doc in ocr_docs}
                docs = [doc for doc in docs if doc.path not in ocr_doc_paths]
                docs = [*docs, *ocr_docs]
                ranked = self._rank_documents(docs, parsed)

        ranked.sort(key=lambda item: item[0], reverse=True)
        results = [self._to_result(doc, score) for score, doc in ranked[:limit]]
        scope = "uploaded files only" if self.uploads_only else "Assignment folder"
        return results, f"{scope}: fast filters first, cached extraction, OCR only when needed"

    def _rank_documents(self, docs: list[LocalDocument], parsed: LocalQuery) -> list[tuple[float, LocalDocument]]:
        if parsed.terms or parsed.phrases:
            semantic = semantic_scores(
                " ".join(parsed.phrases or parsed.terms),
                [f"{doc.path.name}\n{doc.content[:4000]}" for doc in docs],
            )
        else:
            semantic = [0.0 for _doc in docs]

        ranked: list[tuple[float, LocalDocument]] = []
        for index, doc in enumerate(docs):
            score = self._score_document(doc, parsed, semantic[index] if index < len(semantic) else 0.0)
            if score > 0 or (not parsed.terms and not parsed.phrases):
                ranked.append((score, doc))
        return ranked

    def list_uploads(self) -> list[FileResult]:
        self.uploads_dir.mkdir(parents=True, exist_ok=True)
        docs = [LocalDocument(path, "") for path in self._iter_files(uploaded_only=True)]
        return [self._to_result(doc, 0.9) for doc in docs]

    def delete_upload(self, relative_path: str) -> Path:
        candidate = (self.root / relative_path).resolve()
        uploads_root = self.uploads_dir.resolve()
        if uploads_root not in candidate.parents:
            raise ValueError("Only uploaded files can be removed.")
        if not candidate.exists() or not candidate.is_file():
            raise FileNotFoundError("Uploaded file was not found.")
        candidate.unlink()
        return candidate

    def summarize_file(self, path: str, max_chars: int = 700) -> str:
        candidate = Path(path)
        if not candidate.is_absolute():
            candidate = self.root / candidate
        candidate = candidate.resolve()
        if self.root not in candidate.parents and candidate != self.root:
            raise ValueError("File must be inside the configured Assignment folder.")
        doc = self._load_document(candidate, allow_ocr=True)
        if not doc.content:
            return f"{candidate.name} has no extractable text content."
        text = " ".join(doc.content.split())
        return text[:max_chars] + ("..." if len(text) > max_chars else "")

    def save_upload(self, filename: str, content: bytes) -> Path:
        suffix = Path(filename).suffix.lower()
        if suffix not in UPLOAD_EXTENSIONS:
            allowed = ", ".join(sorted(UPLOAD_EXTENSIONS))
            raise ValueError(f"Unsupported file type. Allowed extensions: {allowed}")
        self.uploads_dir.mkdir(parents=True, exist_ok=True)
        safe_name = self._safe_filename(filename)
        destination = self.uploads_dir / safe_name
        counter = 1
        while destination.exists():
            destination = self.uploads_dir / f"{Path(safe_name).stem}-{counter}{suffix}"
            counter += 1
        destination.write_bytes(content)
        return destination

    def _iter_files(self, uploaded_only: bool | None = None) -> list[Path]:
        search_uploads = self.uploads_only if uploaded_only is None else uploaded_only
        base = self.uploads_dir if search_uploads else self.root
        base.mkdir(parents=True, exist_ok=True)
        return sorted(
            path for path in base.rglob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
        )

    def _load_document(self, path: Path, allow_ocr: bool = False) -> LocalDocument:
        cached = self._read_cached_document(path, allow_ocr)
        if cached is not None:
            return LocalDocument(path, cached)

        try:
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                content = self._read_pdf(path, allow_ocr=allow_ocr)
            elif suffix == ".docx":
                content = self._read_docx(path)
            elif suffix == ".txt":
                content = path.read_text(errors="ignore")
            elif suffix == ".csv":
                content = self._read_csv(path)
            elif suffix in {".xlsx", ".xls"}:
                content = self._read_excel(path)
            elif suffix in IMAGE_EXTENSIONS:
                content = self._read_image(path) if allow_ocr else f"Image file: {path.name}"
            elif suffix in VIDEO_EXTENSIONS:
                content = self._read_video(path) if allow_ocr else f"Video file: {path.name}"
            else:
                content = ""
            self._write_cached_document(path, content, allow_ocr)
            return LocalDocument(path, content)
        except Exception:
            logger.exception("Failed to extract local file %s", path)
        return LocalDocument(path, "")

    @staticmethod
    def _read_pdf(path: Path, allow_ocr: bool = False) -> str:
        reader = PdfReader(str(path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if text:
            return text
        if allow_ocr:
            return LocalSearchService._ocr_pdf(path)
        return f"PDF file: {path.name}. No embedded text extracted."

    @staticmethod
    def _read_docx(path: Path) -> str:
        document = Document(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    @staticmethod
    def _read_csv(path: Path) -> str:
        with path.open(newline="", errors="ignore") as handle:
            rows = csv.reader(handle)
            return "\n".join(", ".join(row) for row in rows)

    @staticmethod
    def _read_excel(path: Path) -> str:
        frames = pd.read_excel(path, sheet_name=None)
        return "\n".join(frame.to_csv(index=False) for frame in frames.values())

    @staticmethod
    def _read_image(path: Path) -> str:
        try:
            from PIL import Image
        except ImportError:
            return "Image file. OCR unavailable because Pillow is not installed."

        try:
            text = LocalSearchService._ocr_image(Image.open(path))
        except Exception as exc:
            logger.warning("Image OCR failed for %s: %s", path, exc)
            return "Image file. OCR extraction failed; search can still match filename."
        text = text.strip()
        return f"Image OCR text:\n{text}" if text else "Image file. OCR found no readable text."

    @staticmethod
    def _read_video(path: Path) -> str:
        metadata: dict[str, Any] = {"kind": "video"}
        frame_text: list[str] = []
        try:
            import cv2

            capture = cv2.VideoCapture(str(path))
            frame_count = int(capture.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
            fps = float(capture.get(cv2.CAP_PROP_FPS) or 0)
            duration = round(frame_count / fps, 2) if fps else None
            metadata.update({"frames": frame_count, "fps": round(fps, 2), "duration_seconds": duration})

            if frame_count:
                sample_positions = sorted({0, frame_count // 4, frame_count // 2, (frame_count * 3) // 4, frame_count - 1})
                for position in sample_positions:
                    capture.set(cv2.CAP_PROP_POS_FRAMES, position)
                    ok, frame = capture.read()
                    if ok:
                        text = LocalSearchService._ocr_cv2_frame(frame)
                        if text:
                            frame_text.append(text)
            capture.release()
        except ImportError:
            return "Video file. Install opencv-python-headless for video metadata and frame OCR."
        except Exception as exc:
            logger.warning("Video processing failed for %s: %s", path, exc)
            return "Video file. Video metadata extraction failed; search can still match filename."

        summary = (
            f"Video metadata: duration {metadata.get('duration_seconds')} seconds, "
            f"{metadata.get('frames')} frames, {metadata.get('fps')} fps."
        )
        if frame_text:
            summary += "\nOCR text from sampled frames:\n" + "\n".join(frame_text)
        else:
            summary += " No readable OCR text was found in sampled frames."
        return summary

    @staticmethod
    def _ocr_cv2_frame(frame: Any) -> str:
        try:
            import cv2
            from PIL import Image
        except ImportError:
            return ""

        try:
            rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            return LocalSearchService._ocr_image(Image.fromarray(rgb))
        except Exception:
            return ""

    @staticmethod
    def _ocr_image(image: Any) -> str:
        if not LocalSearchService._tesseract_available():
            return "OCR unavailable because the Tesseract system binary is not installed."

        try:
            import pytesseract
        except ImportError:
            return "OCR unavailable because pytesseract is not installed."

        outputs: list[str] = []
        for candidate in LocalSearchService._ocr_candidates(image):
            for config in (
                "--oem 3 --psm 6",
                "--oem 3 --psm 4",
                "--oem 3 --psm 11",
                "--oem 3 --psm 3",
            ):
                try:
                    text = pytesseract.image_to_string(candidate, config=config).strip()
                except Exception:
                    continue
                if LocalSearchService._ocr_quality(text) >= 8:
                    outputs.append(text)

        return LocalSearchService._merge_ocr_outputs(outputs)

    @staticmethod
    def _ocr_candidates(image: Any) -> list[Any]:
        from PIL import Image, ImageEnhance, ImageFilter, ImageOps

        base = ImageOps.exif_transpose(image).convert("L")
        max_side = max(base.size)
        if max_side < 2200:
            scale = min(4, max(2, int(2400 / max_side)))
            base = base.resize((base.width * scale, base.height * scale), Image.Resampling.LANCZOS)

        contrast = ImageEnhance.Contrast(base).enhance(2.0)
        sharp = contrast.filter(ImageFilter.SHARPEN)
        denoised = sharp.filter(ImageFilter.MedianFilter(size=3))
        autocontrast = ImageOps.autocontrast(denoised)

        candidates = [
            autocontrast,
            autocontrast.point(lambda pixel: 255 if pixel > 165 else 0),
            autocontrast.point(lambda pixel: 255 if pixel > 190 else 0),
        ]

        try:
            import cv2
            import numpy as np

            arr = np.array(autocontrast)
            adaptive = cv2.adaptiveThreshold(
                arr,
                255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY,
                35,
                11,
            )
            candidates.append(Image.fromarray(adaptive))
        except Exception:
            pass

        return candidates

    @staticmethod
    def _ocr_pdf(path: Path, max_pages: int = 3) -> str:
        if not LocalSearchService._tesseract_available():
            return "PDF file. No embedded text found, and OCR unavailable because Tesseract is not installed."

        try:
            from pdf2image import convert_from_path
        except ImportError:
            return "PDF file. No embedded text found. Install pdf2image and poppler to OCR scanned PDFs."

        try:
            pages = convert_from_path(str(path), dpi=180, first_page=1, last_page=max_pages)
        except Exception as exc:
            logger.warning("Scanned PDF OCR conversion failed for %s: %s", path, exc)
            return "PDF file. No embedded text found, and scanned PDF OCR conversion failed."

        page_text = [LocalSearchService._ocr_image(page) for page in pages]
        merged = LocalSearchService._merge_ocr_outputs(page_text)
        return f"Scanned PDF OCR text:\n{merged}" if merged else "PDF file. OCR found no readable text."

    def _read_cached_document(self, path: Path, allow_ocr: bool) -> str | None:
        cache_path = self._cache_path(path, allow_ocr)
        if not cache_path.exists():
            return None
        try:
            payload = json.loads(cache_path.read_text())
            stat = path.stat()
            if payload.get("mtime") == stat.st_mtime and payload.get("size") == stat.st_size:
                return str(payload.get("content", ""))
        except Exception:
            return None
        return None

    def _write_cached_document(self, path: Path, content: str, allow_ocr: bool) -> None:
        try:
            self.cache_dir.mkdir(parents=True, exist_ok=True)
            stat = path.stat()
            payload = {
                "path": str(path),
                "mtime": stat.st_mtime,
                "size": stat.st_size,
                "content": content,
            }
            self._cache_path(path, allow_ocr).write_text(json.dumps(payload))
        except Exception:
            logger.debug("Failed to write extraction cache for %s", path, exc_info=True)

    def _cache_path(self, path: Path, allow_ocr: bool) -> Path:
        mode = "ocr" if allow_ocr else "basic"
        digest = hashlib.sha256(f"{mode}:{path.resolve()}".encode()).hexdigest()
        return self.cache_dir / f"{digest}.json"

    @staticmethod
    def _tesseract_available() -> bool:
        return shutil.which("tesseract") is not None

    @staticmethod
    def _ocr_quality(text: str) -> int:
        return sum(1 for char in text if char.isalnum())

    @staticmethod
    def _merge_ocr_outputs(outputs: list[str]) -> str:
        seen: set[str] = set()
        lines: list[str] = []
        for output in sorted(outputs, key=LocalSearchService._ocr_quality, reverse=True):
            for line in output.splitlines():
                cleaned = re.sub(r"\s+", " ", line).strip()
                if len(cleaned) < 2:
                    continue
                key = cleaned.lower()
                if key not in seen:
                    seen.add(key)
                    lines.append(cleaned)
        return "\n".join(lines)

    def _parse_query(self, query: str) -> LocalQuery:
        lowered = query.lower()
        extensions: set[str] = set()
        filenames = self._filenames(lowered)
        if re.search(r"\bpdfs?\b", lowered):
            extensions.add(".pdf")
        if re.search(r"\b(docx?|word|documents?)\b", lowered):
            extensions.add(".docx")
        if re.search(r"\b(spreadsheets?|sheets?|excel|xlsx?|csv)\b", lowered):
            extensions.update({".xlsx", ".xls", ".csv"})
        if re.search(r"\b(images?|photos?|png|jpe?g|webp|gif)\b", lowered):
            extensions.update(IMAGE_EXTENSIONS)
        if re.search(r"\b(videos?|mp4|mov|m4v)\b", lowered):
            extensions.update(VIDEO_EXTENSIONS)

        modified_after: datetime | None = None
        now = datetime.now(timezone.utc)
        if "last 7 days" in lowered or "past 7 days" in lowered or "last week" in lowered:
            modified_after = now - timedelta(days=7)
        elif "today" in lowered:
            modified_after = now.replace(hour=0, minute=0, second=0, microsecond=0)
        elif "this week" in lowered:
            start = now - timedelta(days=now.weekday())
            modified_after = start.replace(hour=0, minute=0, second=0, microsecond=0)

        phrases = self._phrases(lowered)
        terms = self._terms(lowered)
        return LocalQuery(
            terms=terms,
            phrases=phrases,
            extensions=extensions,
            filenames=filenames,
            modified_after=modified_after,
        )

    @staticmethod
    def _terms(query: str) -> list[str]:
        stopwords = {
            "find",
            "show",
            "search",
            "files",
            "file",
            "about",
            "related",
            "containing",
            "contains",
            "modified",
            "uploaded",
            "created",
            "recently",
            "last",
            "past",
            "days",
            "day",
            "week",
            "this",
            "today",
            "pdf",
            "pdfs",
            "doc",
            "docs",
            "document",
            "documents",
            "spreadsheet",
            "spreadsheets",
            "sheet",
            "sheets",
            "image",
            "images",
            "video",
            "videos",
            "do",
            "ocr",
            "on",
            "and",
            "say",
            "what",
            "it",
            "reads",
            "in",
            "the",
        }
        cleaned = "".join(ch if ch.isalnum() or ch.isspace() else " " for ch in query.lower())
        tokens = [term for term in cleaned.split() if term not in stopwords and not term.isdigit()]
        return tokens

    @staticmethod
    def _phrases(query: str) -> list[str]:
        known_phrases = [
            "machine learning",
            "artificial intelligence",
            "revenue forecast",
            "lead interaction",
            "daily report",
            "curriculum vitae",
        ]
        return [phrase for phrase in known_phrases if phrase in query]

    @staticmethod
    def _filenames(query: str) -> set[str]:
        extensions = "|".join(re.escape(ext.lstrip(".")) for ext in sorted(SUPPORTED_EXTENSIONS, key=len, reverse=True))
        pattern = rf"(?<![\w.-])([\w][\w .-]*?\.({extensions}))(?![\w.-])"
        stop_prefix = re.compile(
            r"^.*\b(?:the|file|called|named|open|show|summarize|read|ocr|does|what|is|in|on)\s+",
            flags=re.IGNORECASE,
        )
        filenames: set[str] = set()
        for match in re.finditer(pattern, query, flags=re.IGNORECASE):
            candidate = stop_prefix.sub("", match.group(1)).strip()
            if candidate:
                filenames.add(LocalSearchService._normalize_filename(candidate))
        return filenames

    def _passes_filters(self, path: Path, parsed: LocalQuery) -> bool:
        if parsed.filenames and not self._filename_matches(path.name, parsed.filenames):
            return False
        if parsed.extensions and path.suffix.lower() not in parsed.extensions:
            return False
        if parsed.modified_after:
            modified = datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.utc)
            if modified < parsed.modified_after:
                return False
        return True

    @staticmethod
    def _filename_matches(path_name: str, requested_filenames: set[str]) -> bool:
        normalized_path = LocalSearchService._normalize_filename(path_name)
        return any(
            normalized_path == requested
            or normalized_path.endswith(requested)
            or requested in normalized_path
            for requested in requested_filenames
        )

    @staticmethod
    def _normalize_filename(filename: str) -> str:
        return re.sub(r"\s+", " ", filename.casefold()).strip()

    @staticmethod
    def _ocr_candidate(path: Path) -> bool:
        return path.suffix.lower() in IMAGE_EXTENSIONS | VIDEO_EXTENSIONS | {".pdf"}

    @staticmethod
    def _should_use_ocr(query: str, parsed: LocalQuery) -> bool:
        lowered = query.lower()
        if any(word in lowered for word in ("ocr", "scanned", "screenshot", "image", "photo", "picture", "video")):
            return True
        if parsed.extensions & (IMAGE_EXTENSIONS | VIDEO_EXTENSIONS):
            return True
        if parsed.extensions == {".pdf"} and any(word in lowered for word in ("text", "content", "contains", "inside")):
            return True
        return False

    @staticmethod
    def _score_document(doc: LocalDocument, parsed: LocalQuery, semantic_score: float) -> float:
        haystack = f"{doc.path.name} {doc.content}".lower()
        filename = doc.path.name.lower()

        phrase_hits = sum(1 for phrase in parsed.phrases if phrase in haystack)
        term_hits = sum(1 for term in parsed.terms if term in haystack)
        filename_hits = sum(1 for term in parsed.terms if term in filename)

        if parsed.phrases and phrase_hits == 0:
            return 0.0
        if not parsed.phrases and parsed.terms and term_hits == 0:
            return 0.0

        phrase_score = 0.55 * phrase_hits / max(len(parsed.phrases), 1) if parsed.phrases else 0.0
        term_score = 0.35 * term_hits / max(len(parsed.terms), 1) if parsed.terms else 0.0
        filename_score = 0.15 * filename_hits / max(len(parsed.terms), 1) if parsed.terms else 0.0
        semantic_component = 0.15 * semantic_score if term_hits or phrase_hits else 0.0
        filter_bonus = 0.1 if parsed.extensions or parsed.modified_after else 0.0
        return round(min(0.98, phrase_score + term_score + filename_score + semantic_component + filter_bonus), 3)

    def _to_result(self, doc: LocalDocument, score: float) -> FileResult:
        stat = doc.path.stat()
        relative = doc.path.relative_to(self.root)
        summary = self._sentence_summary(doc.content) if doc.content else f"Filename match for {doc.path.name}."
        return FileResult(
            name=doc.path.name,
            type=doc.path.suffix.lower().lstrip(".") or "file",
            modified_time=str(pd.Timestamp(stat.st_mtime, unit="s").isoformat()),
            path=str(relative),
            confidence=round(max(0.25, min(0.98, score)), 2),
            source="local",
            summary=summary,
            metadata={"absolute_path": str(doc.path), "size": stat.st_size},
        )

    @staticmethod
    def _safe_filename(filename: str) -> str:
        name = Path(filename).name.strip()
        name = re.sub(r"[^A-Za-z0-9._ -]", "_", name)
        name = re.sub(r"\s+", " ", name)
        return name or "uploaded-file"

    @staticmethod
    def _sentence_summary(text: str, max_chars: int = 420) -> str:
        normalized = " ".join(text.split())
        if len(normalized) <= max_chars:
            return normalized

        sentence_matches = list(re.finditer(r"(?<=[.!?])\s+", normalized))
        best_end = 0
        for match in sentence_matches:
            if match.end() <= max_chars:
                best_end = match.start()
            else:
                break

        if best_end >= 80:
            return normalized[:best_end].strip()

        next_sentence = re.search(r"(?<=[.!?])\s+", normalized[max_chars:])
        if next_sentence:
            end = max_chars + next_sentence.start()
            return normalized[:end].strip()

        return normalized
